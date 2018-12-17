# -*- coding: utf-8 -*-

# Copyright 2018 Wudr. All Rights Reserved.
#
#
# https://github.com/camppolite/ECW.git
# ==============================================================================

from PyQt5.QtWidgets import QMessageBox, QProgressDialog

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn

from ui_file import ui_model2_testreport

class Model2_TestReport():
    """模板2-测试报告的输出"""

    def __init__(self):
        self.model2_testreport_ui = ui_model2_testreport.Ui_Form()

        self.Cases = []

    def ShowFirstRow(self, firstrow):
        """将Sheet的第一行有效列显示，供用户选择"""
        self.model2_testreport_ui.combox_tablename.clear()
        self.model2_testreport_ui.combox_casename.clear()
        self.model2_testreport_ui.combox_caseID.clear()
        self.model2_testreport_ui.combox_requirement.clear()
        self.model2_testreport_ui.combox_condition.clear()
        self.model2_testreport_ui.combox_casetext.clear()
        self.model2_testreport_ui.combox_testdata.clear()
        self.model2_testreport_ui.combox_passcondition.clear()
        self.model2_testreport_ui.combox_input.clear()
        self.model2_testreport_ui.combox_expect.clear()
        self.model2_testreport_ui.combox_result.clear()
        # self.model2_testreport_ui.combox_ispass.clear()
        self.model2_testreport_ui.combox_totalresult.clear()
        self.model2_testreport_ui.combox_bugID.clear()
        self.model2_testreport_ui.combox_testman.clear()
        self.model2_testreport_ui.combox_testtime.clear()

        self.model2_testreport_ui.combox_tablename.addItem("请选择转换列")
        self.model2_testreport_ui.combox_casename.addItem("请选择转换列")
        self.model2_testreport_ui.combox_caseID.addItem("请选择转换列")
        self.model2_testreport_ui.combox_requirement.addItem("请选择转换列")
        self.model2_testreport_ui.combox_condition.addItem("请选择转换列")
        self.model2_testreport_ui.combox_casetext.addItem("请选择转换列")
        self.model2_testreport_ui.combox_testdata.addItem("请选择转换列")
        self.model2_testreport_ui.combox_passcondition.addItem("请选择转换列")
        self.model2_testreport_ui.combox_input.addItem("请选择转换列")
        self.model2_testreport_ui.combox_expect.addItem("请选择转换列")
        self.model2_testreport_ui.combox_result.addItem("请选择转换列")
        # self.model2_testreport_ui.combox_ispass.addItem("请选择转换列")
        self.model2_testreport_ui.combox_totalresult.addItem("请选择转换列")
        self.model2_testreport_ui.combox_bugID.addItem("请选择转换列")
        self.model2_testreport_ui.combox_testman.addItem("请选择转换列")
        self.model2_testreport_ui.combox_testtime.addItem("请选择转换列")

        self.model2_testreport_ui.combox_tablename.addItems(firstrow)
        self.model2_testreport_ui.combox_casename.addItems(firstrow)
        self.model2_testreport_ui.combox_caseID.addItems(firstrow)
        self.model2_testreport_ui.combox_requirement.addItems(firstrow)
        self.model2_testreport_ui.combox_condition.addItems(firstrow)
        self.model2_testreport_ui.combox_casetext.addItems(firstrow)
        self.model2_testreport_ui.combox_testdata.addItems(firstrow)
        self.model2_testreport_ui.combox_passcondition.addItems(firstrow)
        self.model2_testreport_ui.combox_input.addItems(firstrow)
        self.model2_testreport_ui.combox_expect.addItems(firstrow)
        self.model2_testreport_ui.combox_result.addItems(firstrow)
        # self.model2_testreport_ui.combox_ispass.addItems(firstrow)
        self.model2_testreport_ui.combox_totalresult.addItems(firstrow)
        self.model2_testreport_ui.combox_bugID.addItems(firstrow)
        self.model2_testreport_ui.combox_testman.addItems(firstrow)
        self.model2_testreport_ui.combox_testtime.addItems(firstrow)

    def GetComboBoxSelected(self):
        """获取用户的选择列index，返回字典"""
        tablename = self.model2_testreport_ui.combox_tablename.currentIndex()-1
        casename = self.model2_testreport_ui.combox_casename.currentIndex()-1
        caseID = self.model2_testreport_ui.combox_caseID.currentIndex()-1
        requirement = self.model2_testreport_ui.combox_requirement.currentIndex()-1
        condition = self.model2_testreport_ui.combox_condition.currentIndex()-1
        casetext = self.model2_testreport_ui.combox_casetext.currentIndex()-1
        testdata = self.model2_testreport_ui.combox_testdata.currentIndex()-1
        passcondition = self.model2_testreport_ui.combox_passcondition.currentIndex()-1
        input = self.model2_testreport_ui.combox_input.currentIndex()-1
        expect = self.model2_testreport_ui.combox_expect.currentIndex()-1
        result = self.model2_testreport_ui.combox_result.currentIndex()-1
        # ispass = self.model2_testreport_ui.combox_ispass.currentIndex()-1
        totalresult = self.model2_testreport_ui.combox_totalresult.currentIndex()-1
        bugID = self.model2_testreport_ui.combox_bugID.currentIndex()-1
        testman = self.model2_testreport_ui.combox_testman.currentIndex()-1
        testtime = self.model2_testreport_ui.combox_testtime.currentIndex()-1

        ComboBoxSelected = {
            "tablename" : tablename,
            "casename" : casename,
            "caseID" : caseID,
            "requirement" : requirement,
            "condition" : condition,
            "casetext" : casetext,
            "testdata" : testdata,
            "passcondition" : passcondition,
            "input" : input,
            "expect" : expect,
            "result": result,
            # "ispass": ispass,
            "totalresult": totalresult,
            "bugID": bugID,
            "testman" : testman,
            "testtime": testtime
        }

        return ComboBoxSelected

    def CreateWord(self, savePathName):
        """
        将解析的Excel内容转换为Word内容并输出为Word文件
        """

        ComboBoxSelected = self.GetComboBoxSelected()
        tablename = ComboBoxSelected.get("tablename")
        casename = ComboBoxSelected.get("casename")
        caseID = ComboBoxSelected.get("caseID")
        requirement = ComboBoxSelected.get("requirement")
        condition = ComboBoxSelected.get("condition")
        casetext = ComboBoxSelected.get("casetext")
        testdata = ComboBoxSelected.get("testdata")
        passcondition = ComboBoxSelected.get("passcondition")
        input = ComboBoxSelected.get("input")
        expect = ComboBoxSelected.get("expect")
        result = ComboBoxSelected.get("result")
        # ispass = ComboBoxSelected.get("ispass")
        totalresult = ComboBoxSelected.get("totalresult")
        bugID = ComboBoxSelected.get("bugID")
        testman = ComboBoxSelected.get("testman")
        testtime = ComboBoxSelected.get("testtime")

        document = Document()

        # 设置进度显示
        prs = 1
        len_cases = len(self.Cases)
        process = QProgressDialog(self.model2_testreport_ui.groupBox)
        #如果进度条运行的时间小于2秒，进度条就不会显示，不设置默认是4S
        process.setMinimumDuration(2)
        process.setWindowTitle("清稍后")
        process.setLabelText("正在保存...")
        process.setCancelButtonText("取消")
        process.setRange(prs, len_cases+1)
        process.setModal(True)

        for case in self.Cases:
            prs += 1
            process.setValue(prs)

            if process.wasCanceled():
                break

            # 添加表，并设置样式，可以使用的样式参见"site-packages\docx\templates\default-styles.xml"
            # 或"http://python-docx.readthedocs.io/en/latest/user/styles-understanding.html"
            document.add_paragraph(
                case[tablename], style='List Number'
            ).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table = document.add_table(rows=7, cols=6, style='Table Grid')
            #第1行
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = u'测试用例名称'
            hdr_cells[1].merge(hdr_cells[2])
            hdr_cells[1].text = case[casename]
            hdr_cells[3].text = '用例标识'
            hdr_cells[4].merge(hdr_cells[5])
            hdr_cells[4].text = case[caseID]
            #宋体小四加粗
            cell_runs00 = hdr_cells[0].paragraphs[0].runs
            cell_runs00[0].font.name = u'宋体'
            r = cell_runs00[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs00[0].font.size = Pt(12)
            # cell_runs00[0].font.bold = True

            cell_runs02 = hdr_cells[3].paragraphs[0].runs
            cell_runs02[0].font.name = u'宋体'
            r = cell_runs02[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs02[0].font.size = Pt(12)
            # cell_runs04[0].font.bold = True
            # 第2行
            hdr_cells1 = table.rows[1].cells
            hdr_cells1[0].text = '需求追踪'
            hdr_cells1[1].merge(hdr_cells1[2]).merge(hdr_cells1[3]).merge(hdr_cells1[4]).merge(hdr_cells1[5])
            hdr_cells1[1].text = case[requirement]
            #宋体小四加粗
            cell_runs10 = hdr_cells1[0].paragraphs[0].runs
            cell_runs10[0].font.name = u'宋体'
            r = cell_runs10[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs10[0].font.size = Pt(12)
            # cell_runs10[0].font.bold = True

            # cell_runs14 = hdr_cells1[4].paragraphs[0].runs
            # cell_runs14[0].font.name = u'宋体'
            # r = cell_runs14[0]._element
            # r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            # cell_runs14[0].font.size = Pt(12)
            # cell_runs14[0].font.bold = True
            # 第3行
            hdr_cells2 = table.rows[2].cells
            hdr_cells2[0].text = '先决条件'
            hdr_cells2[1].merge(hdr_cells2[2]).merge(hdr_cells2[3]).merge(hdr_cells2[4]).merge(hdr_cells2[5])
            hdr_cells2[1].text = case[condition]
            #宋体小四加粗
            cell_runs20 = hdr_cells2[0].paragraphs[0].runs
            cell_runs20[0].font.name = u'宋体'
            r = cell_runs20[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs20[0].font.size = Pt(12)
            # cell_runs20[0].font.bold = True

            # cell_runs22 = hdr_cells2[2].paragraphs[0].runs
            # cell_runs22[0].font.name = u'宋体'
            # r = cell_runs22[0]._element
            # r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            # cell_runs22[0].font.size = Pt(12)
            # cell_runs22[0].font.bold = True
            #
            # cell_runs24 = hdr_cells2[4].paragraphs[0].runs
            # cell_runs24[0].font.name = u'宋体'
            # r = cell_runs24[0]._element
            # r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            # cell_runs24[0].font.size = Pt(12)
            # cell_runs24[0].font.bold = True
            # 第4行
            hdr_cells3 = table.rows[3].cells
            hdr_cells3[0].text = '测试用例综述'
            hdr_cells3[1].merge(hdr_cells3[2]).merge(hdr_cells3[3]).merge(hdr_cells3[4]).merge(hdr_cells3[5])
            hdr_cells3[1].text = case[casetext]
            #宋体小四加粗
            cell_runs30 = hdr_cells3[0].paragraphs[0].runs
            cell_runs30[0].font.name = u'宋体'
            r = cell_runs30[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs30[0].font.size = Pt(12)
            # cell_runs30[0].font.bold = True
            # 第5行
            hdr_cells4 = table.rows[4].cells
            hdr_cells4[0].text = '测试数据'
            hdr_cells4[1].merge(hdr_cells4[2]).merge(hdr_cells4[3]).merge(hdr_cells4[4]).merge(hdr_cells4[5])
            hdr_cells4[1].text = case[testdata]
            #宋体小四加粗
            cell_runs40 = hdr_cells4[0].paragraphs[0].runs
            cell_runs40[0].font.name = u'宋体'
            r = cell_runs40[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs40[0].font.size = Pt(12)
            # cell_runs40[0].font.bold = True
            # 第6行
            hdr_cells5 = table.rows[5].cells
            hdr_cells5[0].text = '通过准则'
            hdr_cells5[1].merge(hdr_cells5[2]).merge(hdr_cells5[3]).merge(hdr_cells5[4]).merge(hdr_cells5[5])
            hdr_cells5[1].text = case[passcondition]
            #宋体小四加粗
            cell_runs50 = hdr_cells5[0].paragraphs[0].runs
            cell_runs50[0].font.name = u'宋体'
            r = cell_runs50[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs50[0].font.size = Pt(12)
            # cell_runs50[0].font.bold = True
            # 第7行
            hdr_cells6 = table.rows[6].cells
            hdr_cells6[0].merge(hdr_cells6[1]).merge(hdr_cells6[2]).merge(hdr_cells6[3]).\
                merge(hdr_cells6[4]).merge(hdr_cells6[5])
            hdr_cells6[0].text = '测试步骤'
            #宋体小四加粗
            cell_runs60 = hdr_cells6[0].paragraphs[0].runs
            cell_runs60[0].font.name = u'宋体'
            r = cell_runs60[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs60[0].font.size = Pt(12)
            # cell_runs60[0].font.bold = True
            # 第8行
            table2 = document.add_table(rows=2, cols=10, style='Table Grid')

            hdr2_cells = table2.rows[0].cells
            hdr2_cells[0].text = '序号'
            hdr2_cells[1].merge(hdr2_cells[2]).merge(hdr2_cells[3])
            hdr2_cells[1].text = '输入及操作'
            hdr2_cells[4].merge(hdr2_cells[5]).merge(hdr2_cells[6])
            hdr2_cells[4].text = '期望结果'
            hdr2_cells[7].merge(hdr2_cells[8]).merge(hdr2_cells[9])
            hdr2_cells[7].text = '实测结果'
            # hdr2_cells[8].text = '通过与否'
            #宋体小四加粗
            # cell_runs70 = hdr2_cells[0].paragraphs[0].runs
            # cell_runs70[0].font.name = u'宋体'
            # r = cell_runs70[0]._element
            # r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            # cell_runs70[0].font.size = Pt(12)
            # # cell_runs70[0].font.bold = True
            #
            # cell_runs71 = hdr2_cells[1].paragraphs[0].runs
            # cell_runs71[0].font.name = u'宋体'
            # r = cell_runs71[0]._element
            # r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            # cell_runs71[0].font.size = Pt(12)
            #
            # cell_runs73 = hdr2_cells[5].paragraphs[0].runs
            # cell_runs73[0].font.name = u'宋体'
            # r = cell_runs73[0]._element
            # r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            # cell_runs73[0].font.size = Pt(12)

            #第9行
            hdr2_cells1 = table2.rows[1].cells
            hdr2_cells1[0].text = '1'
            hdr2_cells1[1].merge(hdr2_cells1[2]).merge(hdr2_cells1[3])
            hdr2_cells1[1].text = case[input]
            hdr2_cells1[4].merge(hdr2_cells1[5]).merge(hdr2_cells1[6])
            hdr2_cells1[4].text = case[expect]
            hdr2_cells1[7].merge(hdr2_cells1[8]).merge(hdr2_cells1[9])
            hdr2_cells1[7].text = case[result]
            # hdr2_cells1[8].text = case[ispass]

            #宋体小四加粗
            cell_runs80 = hdr2_cells1[0].paragraphs[0].runs
            cell_runs80[0].font.name = u'宋体'
            r = cell_runs80[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs80[0].font.size = Pt(12)
            #第10行
            table3 = document.add_table(rows=4, cols=6, style='Table Grid')

            hdr3_cells = table3.rows[0].cells
            hdr3_cells[0].text = '执行结果'
            hdr3_cells[1].merge(hdr3_cells[2]).merge(hdr3_cells[3]).merge(hdr3_cells[4]).merge(hdr3_cells[5])
            hdr3_cells[1].text = case[totalresult]
            #宋体小四加粗
            cell_runs90 = hdr3_cells[0].paragraphs[0].runs
            cell_runs90[0].font.name = u'宋体'
            r = cell_runs90[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs90[0].font.size = Pt(12)
            #第11行
            hdr3_cells = table3.rows[1].cells
            hdr3_cells[0].text = '问题标识单'
            hdr3_cells[1].merge(hdr3_cells[2]).merge(hdr3_cells[3]).merge(hdr3_cells[4]).merge(hdr3_cells[5])
            hdr3_cells[1].text = case[bugID]
            #宋体小四加粗
            cell_runs100 = hdr3_cells[0].paragraphs[0].runs
            cell_runs100[0].font.name = u'宋体'
            r = cell_runs100[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs100[0].font.size = Pt(12)
            #第12行
            hdr3_cells = table3.rows[2].cells
            hdr3_cells[0].text = '执行人员'
            hdr3_cells[1].merge(hdr3_cells[2]).merge(hdr3_cells[3]).merge(hdr3_cells[4]).merge(hdr3_cells[5])
            hdr3_cells[1].text = case[testman]
            #宋体小四加粗
            cell_runs110 = hdr3_cells[0].paragraphs[0].runs
            cell_runs110[0].font.name = u'宋体'
            r = cell_runs110[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs110[0].font.size = Pt(12)

            #第13行
            hdr3_cells = table3.rows[3].cells
            hdr3_cells[0].text = '测试时间'
            hdr3_cells[1].merge(hdr3_cells[2]).merge(hdr3_cells[3]).merge(hdr3_cells[4]).merge(hdr3_cells[5])
            hdr3_cells[1].text = case[testtime]
            #宋体小四加粗
            cell_runs120 = hdr3_cells[0].paragraphs[0].runs
            cell_runs120[0].font.name = u'宋体'
            r = cell_runs120[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs120[0].font.size = Pt(12)


            # 添加换行段落
            document.add_paragraph()

            #表格文本全部居中
            # for r in range(0, 8):
            #     handr_cells = table.rows[r].cells
            #     for i in range(len(handr_cells)):
            #         for p in range(len(handr_cells[i].paragraphs)):
            #             handr_cells[i].paragraphs[p].alignment = WD_ALIGN_PARAGRAPH.CENTER

            #测试步骤居中显示
            handr_cells = table.rows[6].cells
            handr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        try:
            document.save(savePathName)
        except (IOError, NameError, FileNotFoundError) as e:
            m_box = QMessageBox()
            m_box.setWindowTitle("提示")
            m_box.setText("保存文件失败！！！" + str(e))
            m_box.exec_()
            return

        process.close()
        self.CreateWordSucess()

    def CreateWordSucess(self):
        """弹出转化成功提示框"""
        MsgBox = QMessageBox(parent=self.model2_testreport_ui.groupBox)
        MsgBox.setWindowTitle("提示")
        MsgBox.setText("转换成功！！！")
        MsgBox.exec_()
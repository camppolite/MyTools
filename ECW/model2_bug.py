# -*- coding: utf-8 -*-

# Copyright 2018 Wudr. All Rights Reserved.
#
#
# https://github.com/camppolite/ECW.git
# ==============================================================================

from PyQt5.QtWidgets import QMessageBox, QProgressDialog

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn

from ui_file import ui_model2_bug

class Model2_Bug():
    """模板2-测试报告的输出"""

    def __init__(self):
        self.model2_bug_ui = ui_model2_bug.Ui_Form()

        self.Cases = []

    def ShowFirstRow(self, firstrow):
        """将Sheet的第一行有效列显示，供用户选择"""
        self.model2_bug_ui.combox_tablename.clear()
        self.model2_bug_ui.combox_bugID.clear()
        self.model2_bug_ui.combox_caseID.clear()
        self.model2_bug_ui.combox_projectname.clear()
        self.model2_bug_ui.combox_bugdate.clear()
        self.model2_bug_ui.combox_bugowner.clear()
        self.model2_bug_ui.combox_version.clear()
        self.model2_bug_ui.combox_bugclass.clear()
        self.model2_bug_ui.combox_buglevel.clear()
        self.model2_bug_ui.combox_bugdescrip.clear()
        self.model2_bug_ui.combox_bugexpect.clear()
        self.model2_bug_ui.combox_fixed.clear()
        self.model2_bug_ui.combox_sign.clear()
        self.model2_bug_ui.combox_signdate.clear()

        self.model2_bug_ui.combox_tablename.addItem("请选择转换列")
        self.model2_bug_ui.combox_bugID.addItem("请选择转换列")
        self.model2_bug_ui.combox_caseID.addItem("请选择转换列")
        self.model2_bug_ui.combox_projectname.addItem("请选择转换列")
        self.model2_bug_ui.combox_bugdate.addItem("请选择转换列")
        self.model2_bug_ui.combox_bugowner.addItem("请选择转换列")
        self.model2_bug_ui.combox_version.addItem("请选择转换列")
        self.model2_bug_ui.combox_bugclass.addItem("请选择转换列")
        self.model2_bug_ui.combox_buglevel.addItem("请选择转换列")
        self.model2_bug_ui.combox_bugdescrip.addItem("请选择转换列")
        self.model2_bug_ui.combox_bugexpect.addItem("请选择转换列")
        self.model2_bug_ui.combox_fixed.addItem("请选择转换列")
        self.model2_bug_ui.combox_sign.addItem("请选择转换列")
        self.model2_bug_ui.combox_signdate.addItem("请选择转换列")

        self.model2_bug_ui.combox_tablename.addItems(firstrow)
        self.model2_bug_ui.combox_bugID.addItems(firstrow)
        self.model2_bug_ui.combox_caseID.addItems(firstrow)
        self.model2_bug_ui.combox_projectname.addItems(firstrow)
        self.model2_bug_ui.combox_bugdate.addItems(firstrow)
        self.model2_bug_ui.combox_bugowner.addItems(firstrow)
        self.model2_bug_ui.combox_version.addItems(firstrow)
        self.model2_bug_ui.combox_bugclass.addItems(firstrow)
        self.model2_bug_ui.combox_buglevel.addItems(firstrow)
        self.model2_bug_ui.combox_bugdescrip.addItems(firstrow)
        self.model2_bug_ui.combox_bugexpect.addItems(firstrow)
        self.model2_bug_ui.combox_fixed.addItems(firstrow)
        self.model2_bug_ui.combox_sign.addItems(firstrow)
        self.model2_bug_ui.combox_signdate.addItems(firstrow)

    def GetComboBoxSelected(self):
        """获取用户的选择列index，返回字典"""
        tablename = self.model2_bug_ui.combox_tablename.currentIndex()-1
        bugID = self.model2_bug_ui.combox_bugID.currentIndex()-1
        caseID = self.model2_bug_ui.combox_caseID.currentIndex()-1
        projectname = self.model2_bug_ui.combox_projectname.currentIndex()-1
        bugdate = self.model2_bug_ui.combox_bugdate.currentIndex()-1
        bugowner = self.model2_bug_ui.combox_bugowner.currentIndex()-1
        version = self.model2_bug_ui.combox_version.currentIndex()-1
        bugclass = self.model2_bug_ui.combox_bugclass.currentIndex()-1
        buglevel = self.model2_bug_ui.combox_buglevel.currentIndex()-1
        bugdescrip = self.model2_bug_ui.combox_bugdescrip.currentIndex()-1
        bugexpect = self.model2_bug_ui.combox_bugexpect.currentIndex()-1
        fixed = self.model2_bug_ui.combox_fixed.currentIndex()-1
        sign = self.model2_bug_ui.combox_sign.currentIndex()-1
        signdate = self.model2_bug_ui.combox_signdate.currentIndex()-1

        ComboBoxSelected = {
            "tablename" : tablename,
            "bugID" : bugID,
            "caseID" : caseID,
            "projectname" : projectname,
            "bugdate" : bugdate,
            "bugowner" : bugowner,
            "version" : version,
            "bugclass" : bugclass,
            "buglevel" : buglevel,
            "bugdescrip" : bugdescrip,
            "bugexpect": bugexpect,
            "fixed": fixed,
            "sign": sign,
            "signdate": signdate,
        }

        return ComboBoxSelected

    def CreateWord(self, savePathName):
        """
        将解析的Excel内容转换为Word内容并输出为Word文件
        """

        ComboBoxSelected = self.GetComboBoxSelected()
        tablename = ComboBoxSelected.get("tablename")
        bugID = ComboBoxSelected.get("bugID")
        caseID = ComboBoxSelected.get("caseID")
        projectname = ComboBoxSelected.get("projectname")
        bugdate = ComboBoxSelected.get("bugdate")
        bugowner = ComboBoxSelected.get("bugowner")
        version = ComboBoxSelected.get("version")
        bugclass = ComboBoxSelected.get("bugclass")
        buglevel = ComboBoxSelected.get("buglevel")
        bugdescrip = ComboBoxSelected.get("bugdescrip")
        bugexpect = ComboBoxSelected.get("bugexpect")
        fixed = ComboBoxSelected.get("fixed")
        sign = ComboBoxSelected.get("sign")
        signdate = ComboBoxSelected.get("signdate")

        document = Document()

        # 设置进度显示
        prs = 1
        len_cases = len(self.Cases)
        process = QProgressDialog(self.model2_bug_ui.groupBox)
        #如果进度条运行的时间小于2秒，进度条就不会显示，不设置默认是4S
        process.setMinimumDuration(2)
        process.setWindowTitle("清稍后")
        process.setLabelText("正在保存...")
        process.setCancelButtonText("取消")
        process.setRange(prs, len_cases+1)
        process.setModal(True)

        for case in self.Cases:
            prs += 1
            process.setValue(prs)

            if process.wasCanceled():
                break

            # 添加表，并设置样式，可以使用的样式参见"site-packages\docx\templates\default-styles.xml"
            # 或"http://python-docx.readthedocs.io/en/latest/user/styles-understanding.html"
            document.add_paragraph(
                case[tablename], style='List Number'
            ).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table = document.add_table(rows=4, cols=18, style='Table Grid')
            #第1行
            hdr_cells = table.rows[0].cells
            hdr_cells[0].merge(hdr_cells[1]).merge(hdr_cells[2])
            hdr_cells[0].text = u'问题标识'
            hdr_cells[3].merge(hdr_cells[4]).merge(hdr_cells[5])
            hdr_cells[3].text = case[bugID]
            hdr_cells[6].merge(hdr_cells[7]).merge(hdr_cells[8])
            hdr_cells[6].text = '用例编号'
            hdr_cells[9].merge(hdr_cells[10]).merge(hdr_cells[11])
            hdr_cells[9].text = case[caseID]
            hdr_cells[12].merge(hdr_cells[13]).merge(hdr_cells[14])
            hdr_cells[12].text = "项目名称"
            hdr_cells[15].merge(hdr_cells[16]).merge(hdr_cells[17])
            hdr_cells[15].text = case[projectname]
            #宋体小四加粗
            cell_runs00 = hdr_cells[0].paragraphs[0].runs
            cell_runs00[0].font.name = u'宋体'
            r = cell_runs00[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs00[0].font.size = Pt(12)
            # cell_runs00[0].font.bold = True

            cell_runs06 = hdr_cells[6].paragraphs[0].runs
            cell_runs06[0].font.name = u'宋体'
            r = cell_runs06[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs06[0].font.size = Pt(12)
            # cell_runs04[0].font.bold = True

            cell_runs012 = hdr_cells[12].paragraphs[0].runs
            cell_runs012[0].font.name = u'宋体'
            r = cell_runs012[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs012[0].font.size = Pt(12)
            # cell_runs04[0].font.bold = True
            # 第2行
            hdr_cells1 = table.rows[1].cells
            hdr_cells1[0].merge(hdr_cells1[1]).merge(hdr_cells1[2])
            hdr_cells1[0].text = u'报告日期'
            hdr_cells1[3].merge(hdr_cells1[4]).merge(hdr_cells1[5])
            hdr_cells1[3].text = case[bugdate]
            hdr_cells1[6].merge(hdr_cells1[7]).merge(hdr_cells1[8])
            hdr_cells1[6].text = '报告人'
            hdr_cells1[9].merge(hdr_cells1[10]).merge(hdr_cells1[11])
            hdr_cells1[9].text = case[bugowner]
            hdr_cells1[12].merge(hdr_cells1[13]).merge(hdr_cells1[14])
            hdr_cells1[12].text = "程序/文档名"
            hdr_cells1[15].merge(hdr_cells1[16]).merge(hdr_cells1[17])
            hdr_cells1[15].text = case[version]
            #宋体小四加粗
            cell_runs10 = hdr_cells1[0].paragraphs[0].runs
            cell_runs10[0].font.name = u'宋体'
            r = cell_runs10[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs10[0].font.size = Pt(12)
            # cell_runs00[0].font.bold = True

            cell_runs16 = hdr_cells1[6].paragraphs[0].runs
            cell_runs16[0].font.name = u'宋体'
            r = cell_runs16[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs16[0].font.size = Pt(12)
            # cell_runs04[0].font.bold = True

            cell_runs112 = hdr_cells1[12].paragraphs[0].runs
            cell_runs112[0].font.name = u'宋体'
            r = cell_runs112[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs112[0].font.size = Pt(12)
            # cell_runs04[0].font.bold = True
            # 第3行
            hdr_cells2 = table.rows[2].cells
            hdr_cells2[0].merge(hdr_cells2[1]).merge(hdr_cells2[2])
            hdr_cells2[0].text = '问题类别'
            hdr_cells2[3].merge(hdr_cells2[4]).merge(hdr_cells2[5]).merge(hdr_cells2[6]).merge(hdr_cells2[7])\
                .merge(hdr_cells2[8]).merge(hdr_cells2[9]).merge(hdr_cells2[10]).merge(hdr_cells2[11])\
                .merge(hdr_cells2[12]).merge(hdr_cells2[13]).merge(hdr_cells2[14]).merge(hdr_cells2[15])\
                .merge(hdr_cells2[16]).merge(hdr_cells2[17])

            hdr_cells2[3].text = case[bugclass]
            #宋体小四加粗
            cell_runs20 = hdr_cells2[0].paragraphs[0].runs
            cell_runs20[0].font.name = u'宋体'
            r = cell_runs20[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs20[0].font.size = Pt(12)

            # 第4行
            hdr_cells3 = table.rows[3].cells
            hdr_cells3[0].merge(hdr_cells3[1]).merge(hdr_cells3[2])
            hdr_cells3[0].text = '问题级别'
            hdr_cells3[3].merge(hdr_cells3[4]).merge(hdr_cells3[5]).merge(hdr_cells3[6])\
                .merge(hdr_cells3[7]).merge(hdr_cells3[8]).merge(hdr_cells3[9]).merge(hdr_cells3[10])\
                .merge(hdr_cells3[11]).merge(hdr_cells3[12]).merge(hdr_cells3[13]).merge(hdr_cells3[14])\
                .merge(hdr_cells3[15]).merge(hdr_cells3[16]).merge(hdr_cells3[17])

            hdr_cells3[3].text = case[buglevel]
            #宋体小四加粗
            cell_runs30 = hdr_cells3[0].paragraphs[0].runs
            cell_runs30[0].font.name = u'宋体'
            r = cell_runs30[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs30[0].font.size = Pt(12)

            table2 = document.add_table(rows=3, cols=15, style='Table Grid')

            # 第5行
            hdr_cells4 = table2.rows[0].cells
            hdr_cells4[0].text = '问题描述'
            hdr_cells4[1].merge(hdr_cells4[2]).merge(hdr_cells4[3]).merge(hdr_cells4[4]).merge(hdr_cells4[5])\
                .merge(hdr_cells4[6]).merge(hdr_cells4[7]).merge(hdr_cells4[8]).merge(hdr_cells4[9])\
                .merge(hdr_cells4[10]).merge(hdr_cells4[11]).merge(hdr_cells4[12]).merge(hdr_cells4[13])\
                .merge(hdr_cells4[14])
            hdr_cells4[1].text = case[bugdescrip]
            #宋体小四加粗
            cell_runs40 = hdr_cells4[0].paragraphs[0].runs
            cell_runs40[0].font.name = u'宋体'
            r = cell_runs40[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs40[0].font.size = Pt(12)
            # cell_runs40[0].font.bold = True
            # 第6行
            hdr_cells5 = table2.rows[1].cells
            hdr_cells5[0].text = '修改建议'
            hdr_cells5[1].merge(hdr_cells5[2]).merge(hdr_cells5[3]).merge(hdr_cells5[4]).merge(hdr_cells5[5])\
                .merge(hdr_cells5[6]).merge(hdr_cells5[7]).merge(hdr_cells5[8]).merge(hdr_cells5[9])\
                .merge(hdr_cells5[10]).merge(hdr_cells5[11]).merge(hdr_cells5[12]).merge(hdr_cells5[13])\
                .merge(hdr_cells5[14])
            hdr_cells5[1].text = case[bugexpect]
            #宋体小四加粗
            cell_runs50 = hdr_cells5[0].paragraphs[0].runs
            cell_runs50[0].font.name = u'宋体'
            r = cell_runs50[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs50[0].font.size = Pt(12)
            # 第7行
            hdr_cells6 = table2.rows[2].cells
            hdr_cells6[0].text = '设计师意见'
            hdr_cells6[1].merge(hdr_cells6[2]).merge(hdr_cells6[3]).merge(hdr_cells6[4]).merge(hdr_cells6[5])\
                .merge(hdr_cells6[6]).merge(hdr_cells6[7]).merge(hdr_cells6[8]).merge(hdr_cells6[9])\
                .merge(hdr_cells6[10]).merge(hdr_cells6[11]).merge(hdr_cells6[12]).merge(hdr_cells6[13])\
                .merge(hdr_cells6[14])
            hdr_cells6[1].text = case[fixed]

            #设计师签名和日期右对齐
            hdr_cells6[1].add_paragraph()
            hdr_cells6[1].add_paragraph()
            hdr_cells6[1].add_paragraph()
            hdr_cells6[1].add_paragraph(case[sign]).alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            hdr_cells6[1].add_paragraph(case[signdate]).alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            #宋体小四加粗
            cell_runs60 = hdr_cells6[0].paragraphs[0].runs
            cell_runs60[0].font.name = u'宋体'
            r = cell_runs60[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs60[0].font.size = Pt(12)

            # 添加换行段落
            document.add_paragraph()

            #表格文本全部居中
            # for r in range(0, 8):
            #     handr_cells = table.rows[r].cells
            #     for i in range(len(handr_cells)):
            #         for p in range(len(handr_cells[i].paragraphs)):
            #             handr_cells[i].paragraphs[p].alignment = WD_ALIGN_PARAGRAPH.CENTER

            #测试步骤居中显示
            # handr_cells = table.rows[6].cells
            #             # handr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        try:
            document.save(savePathName)
        except (IOError, NameError, FileNotFoundError) as e:
            m_box = QMessageBox()
            m_box.setWindowTitle("提示")
            m_box.setText("保存文件失败！！！" + str(e))
            m_box.exec_()
            return

        process.close()
        self.CreateWordSucess()

    def CreateWordSucess(self):
        """弹出转化成功提示框"""
        MsgBox = QMessageBox(parent=self.model2_bug_ui.groupBox)
        MsgBox.setWindowTitle("提示")
        MsgBox.setText("转换成功！！！")
        MsgBox.exec_()
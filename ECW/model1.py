# -*- coding: utf-8 -*-

# Copyright 2018 Wudr. All Rights Reserved.
#
#
# https://github.com/camppolite/ECW.git
# ==============================================================================

from PyQt5.QtWidgets import QMessageBox, QProgressDialog

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn

from ui_file import ui_model1

class Model1():
    """模板1的输出"""

    def __init__(self):
        self.model1_ui = ui_model1.Ui_Form()

        self.Cases = []

    def ShowFirstRow(self, firstrow):
        """将Sheet的第一行有效列显示，供用户选择"""
        self.model1_ui.comboBox_tableName.clear()
        self.model1_ui.comboBox_ProjectName.clear()
        self.model1_ui.comboBox_TestNumber.clear()
        self.model1_ui.comboBox_priority.clear()
        self.model1_ui.comboBox_requirementNumber.clear()
        self.model1_ui.comboBox_designer.clear()
        self.model1_ui.comboBox_designTime.clear()
        self.model1_ui.comboBox_requirementModel.clear()
        self.model1_ui.comboBox_TestName.clear()
        self.model1_ui.comboBox_TestType.clear()
        self.model1_ui.comboBox_condition.clear()
        self.model1_ui.comboBox_TestStep.clear()
        self.model1_ui.comboBox_expected.clear()

        self.model1_ui.comboBox_tableName.addItem("请选择转换列")
        self.model1_ui.comboBox_ProjectName.addItem("请选择转换列")
        self.model1_ui.comboBox_TestNumber.addItem("请选择转换列")
        self.model1_ui.comboBox_priority.addItem("请选择转换列")
        self.model1_ui.comboBox_requirementNumber.addItem("请选择转换列")
        self.model1_ui.comboBox_designer.addItem("请选择转换列")
        self.model1_ui.comboBox_designTime.addItem("请选择转换列")
        self.model1_ui.comboBox_requirementModel.addItem("请选择转换列")
        self.model1_ui.comboBox_TestName.addItem("请选择转换列")
        self.model1_ui.comboBox_TestType.addItem("请选择转换列")
        self.model1_ui.comboBox_condition.addItem("请选择转换列")
        self.model1_ui.comboBox_TestStep.addItem("请选择转换列")
        self.model1_ui.comboBox_expected.addItem("请选择转换列")

        self.model1_ui.comboBox_tableName.addItems(firstrow)
        self.model1_ui.comboBox_ProjectName.addItems(firstrow)
        self.model1_ui.comboBox_TestNumber.addItems(firstrow)
        self.model1_ui.comboBox_priority.addItems(firstrow)
        self.model1_ui.comboBox_requirementNumber.addItems(firstrow)
        self.model1_ui.comboBox_designer.addItems(firstrow)
        self.model1_ui.comboBox_designTime.addItems(firstrow)
        self.model1_ui.comboBox_requirementModel.addItems(firstrow)
        self.model1_ui.comboBox_TestName.addItems(firstrow)
        self.model1_ui.comboBox_TestType.addItems(firstrow)
        self.model1_ui.comboBox_condition.addItems(firstrow)
        self.model1_ui.comboBox_TestStep.addItems(firstrow)
        self.model1_ui.comboBox_expected.addItems(firstrow)

    def GetComboBoxSelected(self):
        """获取用户的选择列index，返回字典"""
        tableName = self.model1_ui.comboBox_tableName.currentIndex()-1
        ProjectName = self.model1_ui.comboBox_ProjectName.currentIndex()-1
        TestNumber = self.model1_ui.comboBox_TestNumber.currentIndex()-1
        priority = self.model1_ui.comboBox_priority.currentIndex()-1
        requirementNumber = self.model1_ui.comboBox_requirementNumber.currentIndex()-1
        designer = self.model1_ui.comboBox_designer.currentIndex()-1
        designTime = self.model1_ui.comboBox_designTime.currentIndex()-1
        requirementModel = self.model1_ui.comboBox_requirementModel.currentIndex()-1
        TestName = self.model1_ui.comboBox_TestName.currentIndex()-1
        TestType = self.model1_ui.comboBox_TestType.currentIndex()-1
        condition = self.model1_ui.comboBox_condition.currentIndex()-1
        TestStep = self.model1_ui.comboBox_TestStep.currentIndex()-1
        expected = self.model1_ui.comboBox_expected.currentIndex()-1

        ComboBoxSelected = {
            "tableName": tableName,
            "ProjectName": ProjectName,
            "TestNumber": TestNumber,
            "priority": priority,
            "requirementNumber": requirementNumber,
            "designer": designer,
            "designTime": designTime,
            "requirementModel": requirementModel,
            "TestName": TestName,
            "TestType": TestType,
            "condition": condition,
            "TestStep": TestStep,
            "expected": expected
        }

        return ComboBoxSelected

    def CreateWord(self, savePathName):
        """
        将解析的Excel内容转换为Word内容并输出为Word文件
        """

        ComboBoxSelected = self.GetComboBoxSelected()
        tableName = ComboBoxSelected.get("tableName")
        ProjectName = ComboBoxSelected.get("ProjectName")
        TestNumber = ComboBoxSelected.get("TestNumber")
        priority = ComboBoxSelected.get("priority")
        requirementNumber = ComboBoxSelected.get("requirementNumber")
        designer = ComboBoxSelected.get("designer")
        designTime = ComboBoxSelected.get("designTime")
        requirementModel = ComboBoxSelected.get("requirementModel")
        TestName = ComboBoxSelected.get("TestName")
        TestType = ComboBoxSelected.get("TestType")
        condition = ComboBoxSelected.get("condition")
        TestStep = ComboBoxSelected.get("TestStep")
        expected = ComboBoxSelected.get("expected")

        document = Document()

        # 设置进度显示
        prs = 1
        len_cases = len(self.Cases)
        process = QProgressDialog(self.model1_ui.groupBox)
        #如果进度条运行的时间小于2秒，进度条就不会显示，不设置默认是4S
        process.setMinimumDuration(2)
        process.setWindowTitle("清稍后")
        process.setLabelText("正在保存...")
        process.setCancelButtonText("取消")
        process.setRange(prs, len_cases+1)
        process.setModal(True)

        for case in self.Cases:
            prs += 1
            process.setValue(prs)

            if process.wasCanceled():
                break

            # 添加表，并设置样式，可以使用的样式参见"site-packages\docx\templates\default-styles.xml"
            # 或"http://python-docx.readthedocs.io/en/latest/user/styles-understanding.html"
            document.add_paragraph(
                case[tableName], style='List Number'
            ).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table = document.add_table(rows=8, cols=6, style='Table Grid')
            #第1行
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = u'项目（软件）编号-项目（软件）名称'
            hdr_cells[1].merge(hdr_cells[2]).merge(hdr_cells[3])
            hdr_cells[1].text = case[ProjectName]
            hdr_cells[4].text = '测试用例标识'
            hdr_cells[5].text = case[TestNumber]
            #宋体小四加粗
            cell_runs00 = hdr_cells[0].paragraphs[0].runs
            cell_runs00[0].font.name = u'宋体'
            r = cell_runs00[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs00[0].font.size = Pt(12)
            cell_runs00[0].font.bold = True

            cell_runs04 = hdr_cells[4].paragraphs[0].runs
            cell_runs04[0].font.name = u'宋体'
            r = cell_runs04[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs04[0].font.size = Pt(12)
            cell_runs04[0].font.bold = True
            # 第2行
            hdr_cells1 = table.rows[1].cells
            hdr_cells1[0].text = '优先级'
            hdr_cells1[1].merge(hdr_cells1[2]).merge(hdr_cells1[3])
            hdr_cells1[1].text = case[priority]
            hdr_cells1[4].text = '软件需求/模块标识'
            hdr_cells1[5].text = case[requirementNumber]
            #宋体小四加粗
            cell_runs10 = hdr_cells1[0].paragraphs[0].runs
            cell_runs10[0].font.name = u'宋体'
            r = cell_runs10[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs10[0].font.size = Pt(12)
            cell_runs10[0].font.bold = True

            cell_runs14 = hdr_cells1[4].paragraphs[0].runs
            cell_runs14[0].font.name = u'宋体'
            r = cell_runs14[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs14[0].font.size = Pt(12)
            cell_runs14[0].font.bold = True
            # 第3行
            hdr_cells2 = table.rows[2].cells
            hdr_cells2[0].text = '用例设计人员'
            hdr_cells2[1].text = case[designer]
            hdr_cells2[2].text = '设计时间'
            hdr_cells2[3].text = case[designTime]
            hdr_cells2[4].text = '软件需求/模块名称'
            hdr_cells2[5].text = case[requirementModel]
            #宋体小四加粗
            cell_runs20 = hdr_cells2[0].paragraphs[0].runs
            cell_runs20[0].font.name = u'宋体'
            r = cell_runs20[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs20[0].font.size = Pt(12)
            cell_runs20[0].font.bold = True

            cell_runs22 = hdr_cells2[2].paragraphs[0].runs
            cell_runs22[0].font.name = u'宋体'
            r = cell_runs22[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs22[0].font.size = Pt(12)
            cell_runs22[0].font.bold = True

            cell_runs24 = hdr_cells2[4].paragraphs[0].runs
            cell_runs24[0].font.name = u'宋体'
            r = cell_runs24[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs24[0].font.size = Pt(12)
            cell_runs24[0].font.bold = True
            # 第4行
            hdr_cells3 = table.rows[3].cells
            hdr_cells3[0].text = '用例说明'
            hdr_cells3[1].merge(hdr_cells3[2]).merge(hdr_cells3[3]).merge(hdr_cells3[4]).merge(hdr_cells3[5])
            hdr_cells3[1].text = case[TestName]
            #宋体小四加粗
            cell_runs30 = hdr_cells3[0].paragraphs[0].runs
            cell_runs30[0].font.name = u'宋体'
            r = cell_runs30[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs30[0].font.size = Pt(12)
            cell_runs30[0].font.bold = True
            # 第5行
            hdr_cells4 = table.rows[4].cells
            hdr_cells4[0].text = '测试类型'
            hdr_cells4[1].merge(hdr_cells4[2]).merge(hdr_cells4[3]).merge(hdr_cells4[4]).merge(hdr_cells4[5])
            hdr_cells4[1].text = case[TestType]
            #宋体小四加粗
            cell_runs40 = hdr_cells4[0].paragraphs[0].runs
            cell_runs40[0].font.name = u'宋体'
            r = cell_runs40[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs40[0].font.size = Pt(12)
            cell_runs40[0].font.bold = True
            # 第6行
            hdr_cells5 = table.rows[5].cells
            hdr_cells5[0].text = '前提条件'
            hdr_cells5[1].merge(hdr_cells5[2]).merge(hdr_cells5[3]).merge(hdr_cells5[4]).merge(hdr_cells5[5])
            hdr_cells5[1].text = case[condition]
            #宋体小四加粗
            cell_runs50 = hdr_cells5[0].paragraphs[0].runs
            cell_runs50[0].font.name = u'宋体'
            r = cell_runs50[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs50[0].font.size = Pt(12)
            cell_runs50[0].font.bold = True
            # 第7行
            hdr_cells6 = table.rows[6].cells
            hdr_cells6[0].text = '测试步骤'
            hdr_cells6[1].merge(hdr_cells6[2]).merge(hdr_cells6[3]).merge(hdr_cells6[4]).merge(hdr_cells6[5])
            hdr_cells6[1].text = case[TestStep]
            #宋体小四加粗
            cell_runs60 = hdr_cells6[0].paragraphs[0].runs
            cell_runs60[0].font.name = u'宋体'
            r = cell_runs60[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs60[0].font.size = Pt(12)
            cell_runs60[0].font.bold = True
            # 第8行
            hdr_cells7 = table.rows[7].cells
            hdr_cells7[0].text = '预期输出'
            hdr_cells7[1].merge(hdr_cells7[2]).merge(hdr_cells7[3]).merge(hdr_cells7[4]).merge(hdr_cells7[5])
            hdr_cells7[1].text = case[expected]
            #宋体小四加粗
            cell_runs70 = hdr_cells7[0].paragraphs[0].runs
            cell_runs70[0].font.name = u'宋体'
            r = cell_runs70[0]._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            cell_runs70[0].font.size = Pt(12)
            cell_runs70[0].font.bold = True

            # 添加换行段落
            document.add_paragraph()

            #表格文本全部居中
            for r in range(0, 8):
                handr_cells = table.rows[r].cells
                for i in range(len(handr_cells)):
                    for p in range(len(handr_cells[i].paragraphs)):
                        handr_cells[i].paragraphs[p].alignment = WD_ALIGN_PARAGRAPH.CENTER

        try:
            document.save(savePathName)
        except (IOError, NameError, FileNotFoundError) as e:
            m_box = QMessageBox()
            m_box.setWindowTitle("提示")
            m_box.setText("保存文件失败！！！" + str(e))
            m_box.exec_()
            return

        process.close()
        self.CreateWordSucess()

    def CreateWordSucess(self):
        """弹出转化成功提示框"""
        MsgBox = QMessageBox(parent=self.model1_ui.groupBox)
        MsgBox.setWindowTitle("提示")
        MsgBox.setText("转换成功！！！")
        MsgBox.exec_()